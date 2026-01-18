import html
import re
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from flask import Flask, jsonify, render_template, request, send_file
from markupsafe import Markup, escape
import nltk
from nltk.corpus import wordnet as wn
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

CET4_FILE = Path(__file__).resolve().parent / "wordscheck" / "CET4_words_from_CET46_2016.csv"
CET6_FILE = Path(__file__).resolve().parent / "wordscheck" / "CET4+6_expanded_words_7952.csv"
TOKEN_RE = re.compile(r"[A-Za-z]+(?:[-'][A-Za-z]+)*")
SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}
WORDLIST_EXTENSIONS = {".txt", ".csv"}
NLTK_DATA_DIR = Path(__file__).resolve().parent / "nltk_data"

TRANSLATION_TABLE = str.maketrans(
    {
        "\u2018": "'",
        "\u2019": "'",
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
    }
)

app = Flask(__name__)


def normalize_text(text: str) -> str:
    return text.translate(TRANSLATION_TABLE)


def extract_tokens(text: str) -> List[str]:
    return TOKEN_RE.findall(normalize_text(text))


def normalize_token(token: str) -> str:
    return token.lower()


def normalize_segment_token(segment: str) -> str:
    return normalize_token(normalize_text(segment))


VARIANT_EQUIVALENTS = (
    ("practice", "practise"),
    ("license", "licence"),
    ("defense", "defence"),
    ("offense", "offence"),
    ("pretense", "pretence"),
    ("analyze", "analyse"),
    ("apologize", "apologise"),
    ("organize", "organise"),
    ("recognize", "recognise"),
    ("realize", "realise"),
    ("color", "colour"),
    ("center", "centre"),
    ("meter", "metre"),
    ("liter", "litre"),
    ("kilometer", "kilometre"),
    ("theater", "theatre"),
    ("program", "programme"),
    ("catalog", "catalogue"),
    ("dialog", "dialogue"),
    ("traveler", "traveller"),
    ("jewelry", "jewellery"),
    ("gray", "grey"),
)


def apply_equivalent_variants(words: Set[str]) -> None:
    for first, second in VARIANT_EQUIVALENTS:
        if first in words or second in words:
            words.add(first)
            words.add(second)


IRREGULAR_FORMS = {
    "am": ("be",),
    "is": ("be",),
    "are": ("be",),
    "was": ("be",),
    "were": ("be",),
    "been": ("be",),
    "being": ("be",),
    "does": ("do",),
    "did": ("do",),
    "done": ("do",),
    "has": ("have",),
    "had": ("have",),
    "better": ("good", "well"),
    "best": ("good", "well"),
    "worse": ("bad", "ill"),
    "worst": ("bad", "ill"),
    "farther": ("far",),
    "farthest": ("far",),
    "further": ("far",),
    "furthest": ("far",),
    "less": ("little",),
    "least": ("little",),
    "more": ("many", "much"),
    "most": ("many", "much"),
    "spoke": ("speak",),
    "spoken": ("speak",),
    "wrote": ("write",),
    "written": ("write",),
    "ate": ("eat",),
    "eaten": ("eat",),
    "took": ("take",),
    "taken": ("take",),
    "gave": ("give",),
    "given": ("give",),
    "saw": ("see",),
    "seen": ("see",),
    "sang": ("sing",),
    "sung": ("sing",),
    "went": ("go",),
    "gone": ("go",),
    "knew": ("know",),
    "known": ("know",),
    "flew": ("fly",),
    "flown": ("fly",),
    "drove": ("drive",),
    "driven": ("drive",),
}

CONTRACTION_EXCEPTIONS = {
    "won't": ["will"],
    "can't": ["can"],
    "shan't": ["shall"],
    "ain't": ["am", "is", "are", "has", "have"],
}


def expand_parentheses(word: str) -> List[str]:
    match = re.search(r"\(([^()]*)\)", word)
    if not match:
        return [word]
    start, end = match.span()
    inside = match.group(1)
    before = word[:start]
    after = word[end:]
    variants = []
    for tail in expand_parentheses(after):
        variants.append(before + inside + tail)
        variants.append(before + tail)
    return list({variant for variant in variants if variant})


def apply_suffix_variant(base: str, suffix: str) -> str:
    if not suffix:
        return base
    if base.endswith("ction") and suffix == "xion":
        return base[:-5] + suffix
    if len(suffix) <= len(base):
        return base[: -len(suffix)] + suffix
    return base + suffix


def expand_entry(entry: str) -> Set[str]:
    parts = entry.split("/")
    base_part = parts[0]
    base_variants = expand_parentheses(base_part)
    variants = set(base_variants)
    for part in parts[1:]:
        if part.startswith("G") and base_variants:
            suffix = part[1:]
            for base in base_variants:
                variants.add(apply_suffix_variant(base, suffix))
        else:
            for variant in expand_parentheses(part):
                variants.add(variant)
    return variants


def normalize_word(word: str) -> Optional[str]:
    word = word.strip()
    if not word:
        return None
    word = word.replace("G", "-")
    return word.lower()


def load_word_list(word_file: Path) -> Set[str]:
    words: Set[str] = set()
    with word_file.open(encoding="utf-8") as file:
        for line in file:
            entry = line.strip().lstrip("\ufeff")
            if not entry or entry.lower() == "word":
                continue
            for variant in expand_entry(entry):
                normalized = normalize_word(variant)
                if not normalized:
                    continue
                words.add(normalized)
                if "-" in normalized:
                    words.add(normalized.replace("-", ""))
    return words


def parse_word_list_text(text: str) -> Set[str]:
    words: Set[str] = set()
    for line in text.splitlines():
        entry = line.strip().lstrip("\ufeff")
        if not entry:
            continue
        entry = entry.split(",", 1)[0].strip()
        if not entry or entry.lower() == "word":
            continue
        for variant in expand_entry(entry):
            normalized = normalize_word(variant)
            if not normalized:
                continue
            words.add(normalized)
            if "-" in normalized:
                words.add(normalized.replace("-", ""))
    return words

def build_word_sets() -> Dict[str, Set[str]]:
    sets: Dict[str, Set[str]] = {}
    if CET4_FILE.exists():
        cet4_words = load_word_list(CET4_FILE)
        apply_equivalent_variants(cet4_words)
        sets["cet4"] = cet4_words
    if CET6_FILE.exists():
        cet6_words = load_word_list(CET6_FILE)
        apply_equivalent_variants(cet6_words)
        sets["cet6"] = cet6_words
    return sets


WORD_SETS = build_word_sets()
WORD_LISTS = {level: sorted(words) for level, words in WORD_SETS.items()}
WORD_RANKS = {
    level: {word: index for index, word in enumerate(word_list)}
    for level, word_list in WORD_LISTS.items()
}
DEFAULT_LEVEL = "cet4"
WORDNET_READY = False


def available_levels() -> List[Tuple[str, str]]:
    options: List[Tuple[str, str]] = []
    if "cet4" in WORD_SETS:
        options.append(("cet4", "CET4"))
    if "cet6" in WORD_SETS:
        options.append(("cet6", "CET6 (CET4+6)"))
    options.append(("custom", "Custom word list"))
    return options


def resolve_level(level: Optional[str]) -> str:
    if level in WORD_SETS:
        return level
    if level == "custom":
        return "custom"
    if DEFAULT_LEVEL in WORD_SETS:
        return DEFAULT_LEVEL
    return next(iter(WORD_SETS), "")


def get_word_set(level: str) -> Optional[Set[str]]:
    return WORD_SETS.get(level)


def get_word_rank(level: str) -> Dict[str, int]:
    return WORD_RANKS.get(level, {})


def build_word_rank(word_set: Set[str]) -> Dict[str, int]:
    return {word: index for index, word in enumerate(sorted(word_set))}


def ensure_wordnet() -> bool:
    global WORDNET_READY
    if WORDNET_READY:
        return True
    try:
        if str(NLTK_DATA_DIR) not in nltk.data.path:
            nltk.data.path.append(str(NLTK_DATA_DIR))
        try:
            nltk.data.find("corpora/wordnet")
        except LookupError:
            nltk.download("wordnet", download_dir=str(NLTK_DATA_DIR))
        try:
            nltk.data.find("corpora/omw-1.4")
        except LookupError:
            nltk.download("omw-1.4", download_dir=str(NLTK_DATA_DIR))
        WORDNET_READY = True
        return True
    except Exception:
        return False


def build_custom_word_set(raw_text: str) -> Set[str]:
    words = parse_word_list_text(raw_text)
    apply_equivalent_variants(words)
    return words


def resolve_word_context(
    level: str, custom_words_text: Optional[str]
) -> Tuple[Optional[Set[str]], Dict[str, int]]:
    if level == "custom":
        if not custom_words_text or not custom_words_text.strip():
            return None, {}
        custom_set = build_custom_word_set(custom_words_text)
        if not custom_set:
            return None, {}
        return custom_set, build_word_rank(custom_set)
    word_set = get_word_set(level)
    if not word_set:
        return None, {}
    return word_set, get_word_rank(level)


def decode_uploaded_text(raw: bytes) -> str:
    for encoding in ("utf-8", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_docx_text(raw: bytes) -> str:
    document = Document(BytesIO(raw))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    return "\n".join(paragraphs)


def extract_pdf_text(raw: bytes) -> str:
    reader = PdfReader(BytesIO(raw))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text:
            pages.append(page_text)
    return "\n".join(pages)


def extract_text_from_upload(file_storage) -> str:
    filename = (file_storage.filename or "").lower()
    suffix = Path(filename).suffix
    raw = file_storage.read()
    if suffix == ".txt":
        return decode_uploaded_text(raw)
    if suffix == ".docx":
        return extract_docx_text(raw)
    if suffix == ".pdf":
        return extract_pdf_text(raw)
    return ""


def extract_word_list_from_upload(file_storage) -> str:
    filename = (file_storage.filename or "").lower()
    suffix = Path(filename).suffix
    raw = file_storage.read()
    if suffix in WORDLIST_EXTENSIONS:
        return decode_uploaded_text(raw)
    return ""


def contraction_bases(token: str) -> List[str]:
    if token in CONTRACTION_EXCEPTIONS:
        return CONTRACTION_EXCEPTIONS[token][:]
    bases = []
    if token.endswith("n't") and len(token) > 3:
        bases.append(token[:-3])
    for suffix in ("'re", "'ve", "'ll", "'d", "'m"):
        if token.endswith(suffix):
            bases.append(token[: -len(suffix)])
    if token.endswith("'s"):
        bases.append(token[:-2])
    if token.endswith("'"):
        bases.append(token[:-1])
    return [base for base in bases if base]


def stem_ly(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if token.endswith("ly") and len(token) > 4:
        base = token[:-2]
        candidates.add(base)
        if base.endswith("i") and len(base) > 1:
            candidates.add(base[:-1] + "y")
        candidates.add(base + "le")
    return candidates


def stem_plural(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if len(token) <= 3:
        return candidates
    if token.endswith("ies") and len(token) > 4:
        candidates.add(token[:-3] + "y")
    if token.endswith("ves") and len(token) > 4:
        candidates.add(token[:-3] + "f")
        candidates.add(token[:-3] + "fe")
    if token.endswith("es") and len(token) > 3:
        candidates.add(token[:-2])
    if token.endswith("s") and len(token) > 3:
        candidates.add(token[:-1])
    return candidates


def stem_ing(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if token.endswith("ing") and len(token) > 5:
        base = token[:-3]
        candidates.add(base)
        if len(base) > 1 and base[-1] == base[-2]:
            candidates.add(base[:-1])
        if not base.endswith("e"):
            candidates.add(base + "e")
    return candidates


def stem_ed(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if token.endswith("ed") and len(token) > 4:
        base = token[:-2]
        candidates.add(base)
        if len(base) > 1 and base[-1] == base[-2]:
            candidates.add(base[:-1])
        if base.endswith("i"):
            candidates.add(base[:-1] + "y")
        if not base.endswith("e"):
            candidates.add(base + "e")
    return candidates


def stem_er(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if token.endswith("er") and len(token) > 4:
        base = token[:-2]
        candidates.add(base)
        if base.endswith("i") and len(base) > 1:
            candidates.add(base[:-1] + "y")
        if len(base) > 1 and base[-1] == base[-2]:
            candidates.add(base[:-1])
        if not base.endswith("e"):
            candidates.add(base + "e")
    return candidates


def stem_est(token: str) -> Set[str]:
    candidates: Set[str] = set()
    if token.endswith("est") and len(token) > 5:
        base = token[:-3]
        candidates.add(base)
        if base.endswith("i") and len(base) > 1:
            candidates.add(base[:-1] + "y")
        if len(base) > 1 and base[-1] == base[-2]:
            candidates.add(base[:-1])
        if not base.endswith("e"):
            candidates.add(base + "e")
    return candidates


def generate_candidates(token: str) -> Set[str]:
    candidates: Set[str] = {token}
    if "-" in token:
        candidates.add(token.replace("-", ""))
    if token in IRREGULAR_FORMS:
        candidates.update(IRREGULAR_FORMS[token])
    candidates.update(contraction_bases(token))

    for base in list(candidates):
        if base in IRREGULAR_FORMS:
            candidates.update(IRREGULAR_FORMS[base])
        candidates.update(stem_ly(base))
        candidates.update(stem_plural(base))
        candidates.update(stem_ing(base))
        candidates.update(stem_ed(base))
        candidates.update(stem_er(base))
        candidates.update(stem_est(base))

    return {candidate for candidate in candidates if candidate}


def is_known_word(token: str, word_set: Set[str]) -> bool:
    normalized = normalize_token(token)
    for candidate in generate_candidates(normalized):
        if candidate in word_set:
            return True
    return False


def compute_missing(tokens: List[str], word_set: Set[str]) -> List[str]:
    missing: List[str] = []
    for token in tokens:
        if not is_known_word(token, word_set):
            missing.append(normalize_token(token))
    return missing


def is_missing_segment(segment: str, missing_set: Set[str]) -> bool:
    if not missing_set:
        return False
    return normalize_segment_token(segment) in missing_set


def suggest_replacements(
    word: str, word_set: Set[str], word_rank: Dict[str, int]
) -> List[str]:
    if not word_set or not ensure_wordnet():
        return []
    normalized = normalize_token(word)
    suggestions: Set[str] = set()
    candidates = generate_candidates(normalized)
    for base in candidates:
        if not re.fullmatch(r"[a-z]+", base):
            continue
        try:
            synsets = wn.synsets(base)
        except LookupError:
            return []
        except Exception:
            continue
        for synset in synsets:
            for lemma in synset.lemma_names():
                lemma_normalized = lemma.replace("_", " ").lower()
                if lemma_normalized == normalized:
                    continue
                if not re.fullmatch(r"[a-z- ]+", lemma_normalized):
                    continue
                tokens = [token for token in re.split(r"[\s-]+", lemma_normalized) if token]
                if not tokens:
                    continue
                if len(tokens) == 1:
                    if is_known_word(tokens[0], word_set):
                        suggestions.add(lemma_normalized)
                    continue
                if all(is_known_word(token, word_set) for token in tokens):
                    suggestions.add(lemma_normalized)
    ranked = sorted(
        suggestions,
        key=lambda item: (
            sum(word_rank.get(token, 1_000_000) for token in re.split(r"[\s-]+", item))
            / max(len(re.split(r"[\s-]+", item)), 1),
            len(item),
        ),
    )
    return ranked[:8]


def iter_segments(text: str) -> List[Tuple[str, bool]]:
    normalized = normalize_text(text)
    segments: List[Tuple[str, bool]] = []
    last_index = 0
    for match in TOKEN_RE.finditer(normalized):
        start, end = match.span()
        if start > last_index:
            segments.append((text[last_index:start], False))
        segments.append((text[start:end], True))
        last_index = end
    if last_index < len(text):
        segments.append((text[last_index:], False))
    return segments


def highlight_missing(text: str, missing_set: Set[str]) -> Markup:
    output: List[str] = []
    for segment, is_word in iter_segments(text):
        if is_word and is_missing_segment(segment, missing_set):
            output.append(f'<span class="missing">{escape(segment)}</span>')
        else:
            output.append(str(escape(segment)))
    return Markup("".join(output))


def build_reportlab_markup(text: str, missing_set: Set[str]) -> str:
    parts: List[str] = []
    for segment, is_word in iter_segments(text):
        escaped = html.escape(segment)
        if is_word and is_missing_segment(segment, missing_set):
            parts.append(f'<font backColor="#FFD27D">{escaped}</font>')
        else:
            parts.append(escaped)
    return "".join(parts).replace("\t", "    ")


def build_docx_bytes(text: str, missing_set: Set[str]) -> BytesIO:
    document = Document()
    for line in text.split("\n"):
        paragraph = document.add_paragraph()
        if not line:
            continue
        for segment, is_word in iter_segments(line):
            run = paragraph.add_run(segment)
            if is_word and is_missing_segment(segment, missing_set):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def build_pdf_bytes(text: str, missing_set: Set[str]) -> BytesIO:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    style = ParagraphStyle(
        name="Body",
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )
    story = []
    for line in text.split("\n"):
        if not line.strip():
            story.append(Spacer(1, 12))
            continue
        story.append(Paragraph(build_reportlab_markup(line, missing_set), style))
    document.build(story)
    output.seek(0)
    return output


def build_api_results(
    text: str, word_set: Set[str], include_suggestions: bool, word_rank: Dict[str, int]
) -> Dict[str, object]:
    results = build_results(text, word_set)
    if include_suggestions:
        results["suggestions_enabled"] = True
        if ensure_wordnet():
            for item in results["missing_items"]:
                item["suggestions"] = suggest_replacements(
                    item["word"], word_set, word_rank
                )
            results["suggestions_status"] = "ready"
        else:
            results["suggestions_status"] = "unavailable"
    else:
        results["suggestions_enabled"] = False
        results["suggestions_status"] = "disabled"
    return results


def build_results(text: str, word_set: Set[str]) -> Dict[str, object]:
    tokens = extract_tokens(text)
    missing = compute_missing(tokens, word_set)
    missing_set = set(missing)
    counter = Counter(missing)
    missing_items = sorted(counter.items(), key=lambda item: (-item[1], item[0]))
    missing_rows = [{"word": word, "count": count} for word, count in missing_items]
    return {
        "total_count": len(tokens),
        "missing_count": sum(counter.values()),
        "unique_missing": len(counter),
        "missing_items": missing_rows,
        "unique_list": "\n".join(word for word, _ in missing_items),
        "highlighted": str(highlight_missing(text, missing_set)),
    }


@app.route("/", methods=["GET", "POST"])
def index() -> str:
    input_text = ""
    error: Optional[str] = None
    results: Optional[Dict[str, object]] = None
    selected_level = resolve_level(request.form.get("level") if request.method == "POST" else None)
    custom_words_text = ""
    include_suggestions = request.form.get("suggestions") in {"on", "true"}

    if request.method == "POST":
        file_storage = request.files.get("text_file")
        wordlist_storage = request.files.get("wordlist_file")
        custom_words_text = request.form.get("wordlist_text", "")
        if wordlist_storage and wordlist_storage.filename:
            suffix = Path(wordlist_storage.filename).suffix.lower()
            if suffix not in WORDLIST_EXTENSIONS:
                error = "Unsupported word list type. Please upload a .txt or .csv file."
            else:
                try:
                    custom_words_text = extract_word_list_from_upload(wordlist_storage)
                except Exception:
                    error = "Failed to parse the word list. Please try another file."

        if file_storage and file_storage.filename:
            suffix = Path(file_storage.filename).suffix.lower()
            if suffix not in SUPPORTED_EXTENSIONS:
                error = "Unsupported file type. Please upload a .txt, .docx, or .pdf file."
            else:
                try:
                    input_text = extract_text_from_upload(file_storage)
                except Exception:
                    error = "Failed to parse the file. Please try another file."
        else:
            input_text = request.form.get("text_input", "")

        if not error:
            word_set, word_rank = resolve_word_context(
                selected_level, custom_words_text if selected_level == "custom" else None
            )
            if not input_text.strip():
                error = "Please upload a text file or paste some text."
            elif selected_level == "custom" and not custom_words_text.strip():
                error = "Please upload or paste a custom word list."
            elif not word_set:
                error = "Word list not found for the selected level."
            else:
                results = build_api_results(
                    input_text, word_set, include_suggestions, word_rank
                )

    return render_template(
        "index.html",
        input_text=input_text,
        error=error,
        results=results,
        level_options=available_levels(),
        selected_level=selected_level,
        wordlist_text=custom_words_text,
        suggestions_enabled=include_suggestions,
    )


@app.route("/check", methods=["POST"])
def check() -> tuple[str, int]:
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    if not isinstance(text, str):
        text = ""
    level = resolve_level(data.get("level"))
    include_suggestions = bool(data.get("suggestions"))
    custom_words_text = data.get("custom_words", "")
    word_set, word_rank = resolve_word_context(
        level, custom_words_text if level == "custom" else None
    )
    if not word_set:
        if level == "custom":
            return jsonify({"error": "Please upload or paste a custom word list."}), 400
        return jsonify({"error": "Word list not found for selected level."}), 500

    if not text.strip():
        return (
            jsonify(
                {
                    "total_count": 0,
                    "missing_count": 0,
                    "unique_missing": 0,
                    "missing_items": [],
                    "unique_list": "",
                    "highlighted": str(escape(text)),
                    "suggestions_enabled": include_suggestions,
                    "suggestions_status": "disabled" if not include_suggestions else "empty",
                }
            ),
            200,
        )

    results = build_api_results(text, word_set, include_suggestions, word_rank)
    results["level"] = level
    return jsonify(results), 200


@app.route("/extract", methods=["POST"])
def extract() -> tuple[str, int]:
    file_storage = request.files.get("text_file")
    if not file_storage or not file_storage.filename:
        return jsonify({"error": "No file uploaded."}), 400
    level = resolve_level(request.form.get("level"))
    custom_words_text = request.form.get("custom_words", "")
    include_suggestions = request.form.get("suggestions") in {"true", "on"}
    word_set, word_rank = resolve_word_context(
        level, custom_words_text if level == "custom" else None
    )
    if not word_set:
        if level == "custom":
            return jsonify({"error": "Please upload or paste a custom word list."}), 400
        return jsonify({"error": "Word list not found for selected level."}), 500

    suffix = Path(file_storage.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return jsonify({"error": "Unsupported file type."}), 400

    try:
        text = extract_text_from_upload(file_storage)
    except Exception:
        return jsonify({"error": "Failed to parse the file."}), 400

    results = build_api_results(text, word_set, include_suggestions, word_rank)
    results["text"] = text
    results["level"] = level
    return jsonify(results), 200


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    level = resolve_level(data.get("level"))
    custom_words_text = data.get("custom_words", "")
    word_set, _ = resolve_word_context(
        level, custom_words_text if level == "custom" else None
    )
    if not word_set:
        if level == "custom":
            return jsonify({"error": "Please upload or paste a custom word list."}), 400
        return jsonify({"error": "Word list not found for selected level."}), 500
    if not isinstance(text, str) or not text.strip():
        return jsonify({"error": "No text provided."}), 400

    tokens = extract_tokens(text)
    missing = compute_missing(tokens, word_set)
    output = build_pdf_bytes(text, set(missing))
    return send_file(
        output,
        mimetype="application/pdf",
        as_attachment=True,
        download_name="cet4-highlighted.pdf",
    )


@app.route("/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    level = resolve_level(data.get("level"))
    custom_words_text = data.get("custom_words", "")
    word_set, _ = resolve_word_context(
        level, custom_words_text if level == "custom" else None
    )
    if not word_set:
        if level == "custom":
            return jsonify({"error": "Please upload or paste a custom word list."}), 400
        return jsonify({"error": "Word list not found for selected level."}), 500
    if not isinstance(text, str) or not text.strip():
        return jsonify({"error": "No text provided."}), 400

    tokens = extract_tokens(text)
    missing = compute_missing(tokens, word_set)
    output = build_docx_bytes(text, set(missing))
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="cet4-highlighted.docx",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
